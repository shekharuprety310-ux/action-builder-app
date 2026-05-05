"""Build a company-format scope .docx from New Project - Draft Scope of Work v1.dotx."""
from __future__ import annotations

import base64
import io
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_BREAK


def load_dotx_as_document(dotx_path: Path) -> Document:
    """python-docx rejects .dotx content-type; clone package as .docx in memory."""
    buf = io.BytesIO()
    with zipfile.ZipFile(dotx_path, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    )
                zout.writestr(item, data)
    buf.seek(0)
    return Document(buf)


def _set_paragraph_text(p: Paragraph, text: str) -> None:
    p.text = text


def _add_bold_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _data_url_to_bytes(data_url: str | None) -> bytes | None:
    if not data_url or not isinstance(data_url, str) or "," not in data_url:
        return None
    meta, body = data_url.split(",", 1)
    if ";base64" not in meta:
        return None
    try:
        return base64.b64decode(body)
    except Exception:
        return None


def _style_run(run, *, size: float = 10.0, bold: bool = False, color: str = "1F2937") -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _add_para(
    doc: Document,
    text: str,
    *,
    size: float = 10.0,
    bold: bool = False,
    color: str = "000000",
    before: float = 0,
    after: float = 4,
    line: float = 1.15,
) -> Paragraph:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    _style_run(r, size=size, bold=bold, color=color)
    return p


def _add_separator(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8A8A8A")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_bullet(
    doc: Document,
    text: str,
    *,
    level: int = 0,
    size: float = 10.0,
    color: str = "000000",
    bold: bool = False,
) -> Paragraph:
    p = doc.add_paragraph(style="List Paragraph")
    # Visual nesting for sub-items
    p.paragraph_format.left_indent = Pt(14 * level)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    bullet = "• " if level == 0 else "◦ "
    r = p.add_run(bullet + text)
    _style_run(r, size=size, bold=bold, color=color)
    return p


def _prepend_form_page_image(doc: Document, image_bytes: bytes) -> None:
    """Insert a full-page form image at the very start of the document."""
    p_img = doc.add_paragraph()
    run = p_img.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(6.2))
    p_img.alignment = 1  # center
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)

    body = doc.element.body
    # Move newly created paragraphs to the top (before existing template content).
    body.insert(0, p_break._p)
    body.insert(0, p_img._p)


def _truncate_from_paragraph(doc: Document, start_idx: int) -> None:
    body = doc.element.body
    paras = doc.paragraphs
    if start_idx >= len(paras):
        return
    start_el = paras[start_idx]._element
    to_remove = []
    found = False
    for el in list(body):
        if el == start_el:
            found = True
        if found and el.tag != qn("w:sectPr"):
            to_remove.append(el)
    for el in to_remove:
        body.remove(el)


def find_scope_section_indices(doc: Document) -> tuple[int, int | None, int | None, int | None]:
    """
    Returns (first_cc_paragraph_index, project_para_idx, finish_para_idx, design_value_para_idx).
    """
    title_i = None
    for i, p in enumerate(doc.paragraphs):
        if "Cost Centre Scope – Detailed Preliminary Draft" in (p.text or ""):
            title_i = i
            break
    if title_i is None:
        return 72, None, None, None

    proj_i = ass_i = dc_val_i = None
    for j in range(title_i + 1, min(title_i + 12, len(doc.paragraphs))):
        raw = doc.paragraphs[j].text or ""
        if raw.startswith("Project:"):
            proj_i = j
        elif raw.startswith("Assumed Finish Level"):
            ass_i = j
        elif "Design Characteristics (from plans):" in raw and j + 1 < len(doc.paragraphs):
            dc_val_i = j + 1

    start_scan = (dc_val_i + 1) if dc_val_i is not None else title_i + 4
    cc_i = None
    for j in range(start_scan, min(start_scan + 8, len(doc.paragraphs))):
        t = (doc.paragraphs[j].text or "").strip()
        if t:
            cc_i = j
            break
    if cc_i is None:
        cc_i = 72
    return cc_i, proj_i, ass_i, dc_val_i


def build_scope_docx(template_path: Path, payload: dict[str, Any]) -> bytes:
    """
    payload keys:
      project: client, name, location, date, num, finishLevel, designChar, archDrawings, structDrawings, selectedRooms[]
      costCentres: [{ name, projectCategory, descriptionNewBuild, descriptionRenovation, items[], exclusions[], notes[] }]
    """
    doc = load_dotx_as_document(template_path)
    proj = payload.get("project") or {}
    centres = payload.get("costCentres") or []

    client = (proj.get("client") or "").strip()
    name = (proj.get("name") or "").strip()
    location = (proj.get("location") or "").strip()
    date = (proj.get("date") or "").strip()
    ref = (proj.get("num") or "").strip()
    finish = (proj.get("finishLevel") or "Medium–High").strip()
    design = (proj.get("designChar") or "—").strip()
    is_new_build = "new build" in str(proj.get("type") or "").lower()
    start_form_img = _data_url_to_bytes(proj.get("startFormImage"))
    ref_form_img = _data_url_to_bytes(proj.get("referenceFormImage"))
    project_photo_img = _data_url_to_bytes(proj.get("projectPhoto"))
    rooms = proj.get("selectedRooms") or []
    rooms_txt = ", ".join(str(x) for x in rooms) if rooms else "—"

    # Header value lines (indices from template)
    if len(doc.paragraphs) > 18:
        _set_paragraph_text(doc.paragraphs[5], "\t\t\t\t\t\t" + client)
        _set_paragraph_text(doc.paragraphs[7], "\t\t" + name)
        _set_paragraph_text(doc.paragraphs[9], "\t\t" + location)
        # Important: Document Date must only use project date, not project number.
        _set_paragraph_text(doc.paragraphs[17], "\t\t" + date)
        _set_paragraph_text(doc.paragraphs[18], "")

    arch = (proj.get("archDrawings") or "").strip()
    struct = (proj.get("structDrawings") or "").strip()
    if arch and len(doc.paragraphs) > 26:
        _set_paragraph_text(doc.paragraphs[26], f"Dated {date}" if date else "Dated ")
    if struct and len(doc.paragraphs) > 32:
        _set_paragraph_text(doc.paragraphs[32], f"Dated {date}" if date else "Dated ")

    cc_start, proj_i, ass_i, dc_val_i = find_scope_section_indices(doc)
    if proj_i is not None:
        _set_paragraph_text(doc.paragraphs[proj_i], f"Project: \n{name}\n")
    if ass_i is not None:
        _set_paragraph_text(doc.paragraphs[ass_i], f"Assumed Finish Level: {finish}\n")
    if dc_val_i is not None:
        _set_paragraph_text(doc.paragraphs[dc_val_i], design)

    _truncate_from_paragraph(doc, cc_start)

    # Front pages: if provided, they must appear first (not buried later).
    # Prepend in reverse order so start-form remains the very first page.
    if ref_form_img:
        _prepend_form_page_image(doc, ref_form_img)
    if start_form_img:
        _prepend_form_page_image(doc, start_form_img)
    elif project_photo_img:
        # If no custom first form is provided, use project picture as the first visual page.
        _prepend_form_page_image(doc, project_photo_img)

    for idx, block in enumerate(centres):
        cc_name = (block.get("name") or "").strip()
        if not cc_name:
            continue
        desc_nb = (block.get("descriptionNewBuild") or "").strip()
        desc_re = (block.get("descriptionRenovation") or "").strip()
        items = block.get("items") or []
        exclusions = block.get("exclusions") or []
        notes = block.get("notes") or []

        _add_para(doc, cc_name, size=11, bold=True, color="000000", after=6)
        if is_new_build:
            if desc_nb:
                _add_para(doc, "New Build", size=10.5, bold=True, color="000000", after=2)
                _add_para(doc, desc_nb, size=10, color="000000", after=5)
        else:
            if desc_re:
                _add_para(doc, "Renovation", size=10.5, bold=True, color="000000", after=2)
                _add_para(doc, desc_re, size=10, color="000000", after=5)
        _add_para(doc, "Assumptions:", size=10, bold=True, color="000000", after=2)
        _add_para(doc, "Scope items:", size=10, bold=True, color="000000", after=3)
        for it in items:
            item_type = str(it.get("type") or "").strip().lower()
            if item_type == "subheading":
                _add_para(doc, str(it.get("name") or ""), size=10.2, bold=True, color="000000", before=4, after=2)
                continue
            line = (it.get("name") or "").strip()
            rooms_list = it.get("rooms") if isinstance(it.get("rooms"), list) else None
            if rooms_list:
                rooms_list = [str(r).strip() for r in rooms_list if str(r or "").strip()]
            room_label = ", ".join(rooms_list) if rooms_list else (it.get("room") or "").strip()
            qty = it.get("qty")
            prod = (it.get("productLine") or "").strip()
            if room_label:
                line += f" ({room_label})"
            if qty not in (None, "", 1, "1"):
                line += f" — Qty: {qty}"
            _add_bullet(doc, line, level=0, size=10, color="000000")
            # If item has multiple product selections, render nested bullets.
            if prod:
                parts = [x.strip() for x in str(prod).split(";") if x.strip()]
                if not parts:
                    parts = [prod]
                for p_line in parts:
                    _add_bullet(doc, p_line, level=1, size=9.5, color="000000")
        if exclusions:
            _add_para(doc, "Exclusions", size=10.5, bold=True, color="000000", before=6, after=2)
            _add_para(doc, "The following items will not be allowed for in our initial pricing model:", size=10, color="000000", after=3)
            for ex in exclusions:
                _add_bullet(doc, str(ex), level=0, size=10, color="000000")
        if notes:
            _add_para(doc, "NOTES:", size=10.5, bold=True, color="000000", before=6, after=2)
            for n in notes:
                _add_bullet(doc, str(n), level=0, size=10, color="000000")
        if idx < len(centres) - 1:
            _add_separator(doc)

    _add_para(doc, f"Selected rooms (Scope Builder): {rooms_txt}", size=9.5, bold=False, color="000000", before=8, after=0)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
